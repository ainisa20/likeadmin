#!/usr/bin/env python3
import json
import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
STRUCTURED_JSON = os.path.join(BASE_DIR, 'fragments', 'fragments_structured.json')

with open(STRUCTURED_JSON, 'r', encoding='utf-8') as f:
    fragments = json.load(f)

doc = Document()

title = doc.add_heading('开办企业一窗通 操作指南知识库', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_intro = doc.add_paragraph()
r1 = p_intro.add_run('以下内容按操作步骤拆分，每个片段对应一张界面截图和操作说明。')
r1.font.size = Pt(11)

p_intro2 = doc.add_paragraph()
r2 = p_intro2.add_run('每个片段含匹配关键词(keywords)、步骤标题(title)、操作说明(context_text)、截图(image)、实操经验提示(tips)。')
r2.font.size = Pt(10)
r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

embedded = 0
failed = 0

for frag in fragments:
    fid = frag['id']
    page = frag['page']
    section = frag.get('section', '') or '\u5176\u4ed6'
    title_text = frag.get('title', '')
    step_label = frag.get('step_label', '')
    keywords = frag.get('keywords', [])
    context_text = frag.get('context_text', '')
    full_page_text = frag.get('full_page_text', '')
    image_file = frag.get('image_file')
    image_exists = frag.get('image_exists', False)
    tips = frag.get('tips', '')

    doc.add_heading(f'## \u7247\u6bb5{fid}\uff08\u7b2c{page}\u9875 \u00b7 {section}\uff09', level=2)

    if step_label:
        p_step = doc.add_paragraph()
        r_step = p_step.add_run(f'\u6b65\u9aa4\uff1a{step_label}')
        r_step.bold = True
        r_step.font.size = Pt(10)
        r_step.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

    if title_text:
        p_title = doc.add_paragraph()
        r_t = p_title.add_run(f'\u6807\u9898\uff1a{title_text}')
        r_t.bold = True
        r_t.font.size = Pt(10.5)

    if keywords:
        p_kw = doc.add_paragraph()
        r_kw_label = p_kw.add_run('\u5173\u952e\u8bcd\uff1a')
        r_kw_label.bold = True
        r_kw_label.font.size = Pt(9)
        r_kw_label.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        r_kw_text = p_kw.add_run(' | '.join(keywords))
        r_kw_text.font.size = Pt(9)
        r_kw_text.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if context_text:
        p = doc.add_paragraph()
        p.add_run(context_text).font.size = Pt(11)

    if image_file and image_exists:
        img_path = os.path.join(BASE_DIR, image_file)
        if os.path.exists(img_path):
            doc.add_paragraph('')
            try:
                img = Image.open(img_path)
                buf = io.BytesIO()
                img.save(buf, format='PNG')
                buf.seek(0)
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = p_img.add_run()
                run_img.add_picture(buf, width=Inches(5.5))
                cap = doc.add_paragraph(f'[\u622a\u56fe\uff1a\u7247\u6bb5{fid}]')
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].font.size = Pt(9)
                cap.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                embedded += 1
            except Exception as e:
                err = doc.add_paragraph(f'[\u56fe\u7247\u65e0\u6cd5\u5d4c\u5165\uff1a{image_file}]')
                err.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                failed += 1

    if full_page_text and full_page_text != context_text:
        p_ctx = doc.add_paragraph()
        rl = p_ctx.add_run('\u3010\u540c\u9875\u4e0a\u4e0b\u6587\u3011')
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        rt = p_ctx.add_run(full_page_text)
        rt.font.size = Pt(9)
        rt.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    if tips:
        doc.add_paragraph('')
        p_tip = doc.add_paragraph()
        r_tip_label = p_tip.add_run('\U0001f4a1 \u5b9e\u64cd\u7ecf\u9a8c\uff1a')
        r_tip_label.bold = True
        r_tip_label.font.size = Pt(10.5)
        r_tip_label.font.color.rgb = RGBColor(0x16, 0x6B, 0x3A)
        r_tip_text = p_tip.add_run(tips)
        r_tip_text.font.size = Pt(10.5)
        r_tip_text.font.color.rgb = RGBColor(0x16, 0x6B, 0x3A)

    doc.add_paragraph('\u2500' * 60)

output_path = os.path.join(BASE_DIR, '\u5f00\u529e\u4f01\u4e1a\u4e00\u7a97\u901a_\u64cd\u4f5c\u6307\u5357\u77e5\u8bc6\u5e93.docx')
doc.save(output_path)
print(f'Generated: {output_path}')
print(f'Total fragments: {len(fragments)}')
print(f'Images embedded: {embedded}')
print(f'Images failed: {failed}')
tips_count = sum(1 for f in fragments if f.get('tips'))
print(f'Fragments with tips: {tips_count}')
